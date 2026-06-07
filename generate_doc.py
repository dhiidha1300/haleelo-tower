#!/usr/bin/env python3
"""Generate Haleelo Tower Platform Implementation Plan v2.0 Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Colors
NAVY = RGBColor(31, 73, 125)
STEEL = RGBColor(70, 130, 180)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = RGBColor(220, 230, 242)
LIGHT_GRAY = RGBColor(242, 242, 242)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    """Set table cell background colour. rgb is a tuple (r,g,b) or RGBColor."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    r, g, b = rgb[0], rgb[1], rgb[2]
    hex_color = '{:02X}{:02X}{:02X}'.format(r, g, b)
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Add thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = STEEL
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = STEEL
    run.font.name = 'Calibri'
    return p

def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def make_table(doc, headers, rows, col_widths=None):
    """Create a styled table with navy headers and alternating rows."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_borders(table)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = LIGHT_BLUE if r_idx % 2 == 0 else RGBColor(255, 255, 255)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("HALEELO TOWER PLATFORM")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Implementation Plan & Feature Specification")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = STEEL
run.font.name = 'Calibri'

doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run("v2.0 — Updated with Accounting, Booking Flow & Product Management")
run.font.size = Pt(12)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("Date: 7 June 2026")
run.font.size = Pt(11)
run.font.name = 'Calibri'

doc.add_paragraph()

prep_p = doc.add_paragraph()
prep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = prep_p.add_run("Prepared following 3 rounds of discovery sessions\n— accounting questionnaire session June 2026")
run.font.size = Pt(11)
run.font.name = 'Calibri'
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf_p.add_run("CONFIDENTIAL")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Executive Summary", level=1)

add_body(doc,
    "This document is Version 2.0 of the Haleelo Tower Platform Implementation Plan. It supersedes "
    "v1.1 and incorporates the findings from the Finance Officer's requirements interview and client "
    "briefing sessions conducted in June 2026."
)
add_body(doc,
    "The following major areas have been added or updated in this version:"
)
add_bullet(doc, "Finance Officer's full accounting requirements, including a complete double-entry accounting module aligned with professional accounting standards")
add_bullet(doc, "Updated conference hall booking approval chain: a 4-step flow (Draft → Admin Pending → Accountant Pending → Booking Approved), with the Finance Officer serving as the final approver before a booking is confirmed")
add_bullet(doc, "Catering packages (Silver, Gold, Platinum) with configurable included services, pricing, DJ and Cameraman add-ons, selectable at booking time")
add_bullet(doc, "Product Management module: all building offerings are managed as Products with dedicated pages, pricing, and associated add-on services")
add_bullet(doc, "Vendor Purchase Orders and Vendor Bills module for procurement of food, drinks, cleaning, and event materials")
add_bullet(doc, "Account Management for the building's 5 separate accounts (Edahab and ZAAD split by purpose, plus Darasalam Bank as the main account)")
add_bullet(doc, "Bill Code & Account Code System: every invoice and expense carries a unique system-generated reference code traceable to the originating document")
add_bullet(doc, "Electricity Bill Management: admin records meter readings, rates are configured in system settings, and charges are auto-calculated and billed per tenant")
add_bullet(doc, "Chart of Accounts with unique numeric codes across Assets, Liabilities, Revenue, and Expenses")
add_bullet(doc, "General Journal / Activity Log: a chronological double-entry ledger of all financial events")
add_bullet(doc, "Trial Balance report: confirms all accounts balance at any given date")
add_bullet(doc, "Central System Configuration module accessible by Super Admin for all platform-wide settings")
add_body(doc,
    "Implementation will proceed across 5 phases, totalling approximately 8 weeks of development, "
    "followed by end-to-end testing, staff training, and go-live. The platform will be deployed "
    "across three subdomains on a VPS environment."
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. BUILDING & SPACE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Building & Space Overview", level=1)
add_body(doc, "Haleelo Tower is a multi-floor commercial building located in Somalia. The building offers a range of rentable spaces including individual office rooms, conference halls of varying capacities, and an educational facility leased to a university tenant.")

add_heading(doc, "2.1 Floor-by-Floor Breakdown", level=2)

floor_headers = ["Floor", "Type / Use", "Space Details"]
floor_rows = [
    ["Basement",      "Educational Facility",   "Entire floor leased to a university tenant as an educational facility; semester-based lease"],
    ["Ground Floor",  "Individual Office Rooms", "10 individual office rooms available for monthly lease"],
    ["1st Floor",     "Offices + Conference Hall", "3 office rooms + 1 conference hall (50-seat capacity)"],
    ["2nd Floor",     "Offices + Conference Hall", "2 office rooms + 1 conference hall (100-seat capacity)"],
    ["3rd Floor",     "Main Conference Hall",    "1 main conference hall (500-seat capacity)"],
]
make_table(doc, floor_headers, floor_rows, col_widths=[3.5, 4.5, 8.5])

add_heading(doc, "2.2 Space Categories & Booking Types", level=2)

cat_headers = ["Space Category", "Booking Type", "Billing Cycle", "Notes"]
cat_rows = [
    ["Conference Hall (50-seat)",   "Per session / custom duration", "Per booking",    "1st Floor; 3 standard sessions per day"],
    ["Conference Hall (100-seat)",  "Per session / custom duration", "Per booking",    "2nd Floor; 3 standard sessions per day"],
    ["Conference Hall (500-seat)",  "Per session / custom duration", "Per booking",    "3rd Floor; 3 standard sessions per day"],
    ["Office Rooms (Ground Floor)", "Monthly lease",                 "Monthly",        "10 individual rooms"],
    ["Office Rooms (1st Floor)",    "Monthly lease",                 "Monthly",        "3 rooms"],
    ["Office Rooms (2nd Floor)",    "Monthly lease",                 "Monthly",        "2 rooms"],
    ["Educational Facility",        "Semester lease",                "Per semester",   "Basement; whole floor to university tenant"],
]
make_table(doc, cat_headers, cat_rows, col_widths=[4.5, 3.5, 2.8, 5.7])

add_heading(doc, "2.3 Conference Hall Session Times", level=2)

session_headers = ["Session", "Start Time", "End Time", "Duration"]
session_rows = [
    ["Morning",   "8:00 AM",  "1:00 PM",  "5 hours"],
    ["Afternoon", "3:00 PM",  "6:30 PM",  "3.5 hours"],
    ["Evening",   "7:00 PM",  "11:00 PM", "4 hours"],
]
make_table(doc, session_headers, session_rows, col_widths=[4, 3.5, 3.5, 5.5])

add_body(doc, "")
add_body(doc, "Custom duration bookings are also supported: minimum 1 hour, maximum 1 month. Recurring bookings (e.g. every Monday morning) are supported. Session times are configurable in the Central System Settings.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PLATFORM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Platform Architecture", level=1)
add_body(doc, "The platform operates across three distinct subdomains, each serving a different audience with a tailored interface and set of permissions.")

arch_headers = ["Subdomain", "Audience", "Purpose"]
arch_rows = [
    ["halelotower.so",        "General Public",  "Public marketing website — building showcase, space listings, online booking request form, product pages"],
    ["portal.halelotower.so", "Tenants",         "Secure tenant self-service portal — view invoices, submit maintenance requests, download lease documents, track bookings"],
    ["admin.halelotower.so",  "Staff",           "Full internal management dashboard — all modules: bookings, finance, accounting, HR, communications, settings"],
]
make_table(doc, arch_headers, arch_rows, col_widths=[5, 3.5, 8])

add_body(doc, "")
add_body(doc, "Technology Stack summary: Next.js (frontend), Laravel (backend API), PostgreSQL (database), AWS S3 (document and photo storage), WhatsApp Business API (automated messaging), SMTP/SendGrid (email), VPS hosting with local-first testing environment.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. User Roles & Permissions", level=1)
add_body(doc, "The platform has 4 internal staff roles and 2 external user types. Role-based access control is enforced at the API level — not just in the UI — so no role can bypass its permissions. The Finance Officer holds a critical position as the final approver in the conference hall booking chain.")

add_heading(doc, "4.1 Role Definitions", level=2)

role_headers = ["Role", "Who", "Description"]
role_rows = [
    ["Super Admin",              "IT / System Administrator (1 person)", "Full system access including system settings, chart of accounts, user management, audit logs. Responsible for platform configuration and security."],
    ["Admin / Building Manager", "Building Manager",                    "Day-to-day operational management: first-level approval of conference bookings, tenant management, lease oversight, maintenance coordination. Cannot edit financial accounts."],
    ["Operations / Receptionist","Operations Staff (×2)",               "Front-line staff: create bookings on behalf of clients, onboard new tenants, generate invoices for Finance review. Cannot access financial reports or accounting module."],
    ["Finance / Accountant",     "Finance Officer",                     "Sole responsibility for all financial operations: invoices, payments, expenses, vendor bills, payroll, chart of accounts, general journal, reports. Final approver for all conference hall bookings."],
    ["Tenant",                   "Building tenants (external)",         "Access to tenant portal only: view own invoices, submit maintenance requests, download documents. No access to admin dashboard."],
    ["Guest / Public",           "General public (external)",           "Access to public website only: view space information, submit online booking request form."],
]
make_table(doc, role_headers, role_rows, col_widths=[4, 4, 8.5])

add_heading(doc, "4.2 Permissions Matrix", level=2)
add_body(doc, "The table below shows the access level each role has across each major platform module. Legend: Full = full read/write/delete; Limited = read and specific actions only; Read = view only; None = no access.")

perm_headers = ["Module", "Super Admin", "Admin", "Operations", "Finance"]
perm_rows = [
    ["Bookings — View",                "Full",    "Full",    "Full",    "Full"],
    ["Bookings — Create (on behalf)",  "Full",    "Full",    "Full",    "None"],
    ["Bookings — Admin Approval",      "Full",    "Full",    "None",    "None"],
    ["Bookings — Finance Approval",    "Full",    "None",    "None",    "Full"],
    ["Bookings — Cancel / Reschedule", "Full",    "Full",    "Limited", "None"],
    ["Tenant Management",              "Full",    "Full",    "Full",    "Read"],
    ["Lease Management",               "Full",    "Full",    "Limited", "Read"],
    ["Product Management",             "Full",    "Full",    "None",    "None"],
    ["Invoices — Create / Send",       "Full",    "None",    "Full",    "Full"],
    ["Invoices — View",                "Full",    "Full",    "Full",    "Full"],
    ["Payments — Record",              "Full",    "None",    "None",    "Full"],
    ["Vendor Bills / Procurement",     "Full",    "None",    "None",    "Full"],
    ["Electricity Billing",            "Full",    "Full",    "None",    "Full"],
    ["Chart of Accounts",              "Full",    "None",    "None",    "Full"],
    ["Account Management",             "Full",    "None",    "None",    "Full"],
    ["Journal / General Ledger",       "Full",    "None",    "None",    "Full"],
    ["Financial Reports",              "Full",    "None",    "None",    "Full"],
    ["HR / Payroll",                   "Full",    "None",    "None",    "Full"],
    ["Maintenance Requests",           "Full",    "Full",    "Full",    "None"],
    ["Communications / Broadcasts",    "Full",    "Full",    "Limited", "None"],
    ["System Settings",                "Full",    "Limited", "None",    "None"],
    ["User Management",                "Full",    "Limited", "None",    "None"],
    ["Audit Trail",                    "Full",    "Read",    "None",    "Read"],
]
make_table(doc, perm_headers, perm_rows, col_widths=[6, 2.8, 2.8, 2.8, 2.8])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PHASE 1 — USER MANAGEMENT, ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Phase 1 — User Management, Roles & Permissions (Pilot)", level=1)
add_body(doc, "Phase 1 establishes the foundational infrastructure for the entire platform: the authentication system, all user accounts and roles, the audit trail, and the central configuration panel that all subsequent phases depend on.")

add_heading(doc, "5.1 Central System Settings & Configuration", level=2)
add_body(doc, "A single unified settings area in the admin dashboard, accessible to Super Admin (full access) and Admin (limited access). All platform-wide parameters are managed here. This module must be completed before any other module is configured.")

settings_areas = [
    ("General Settings", "Building name, logo, contact details, physical address"),
    ("Electricity Rate Configuration", "Set the price per kWh in USD. Rates are versioned with date-stamps: historical readings always use the rate that was active at the time of the reading, ensuring accurate retrospective billing"),
    ("Catering Package Configuration", "Create and edit the three catering packages (Silver, Gold, Platinum). Per package: name, list of included services (e.g. buffet for N people, drinks, waitstaff), and package price. Also configure DJ option (yes/no + price) and Cameraman option (yes/no + price)"),
    ("Session Time Configuration", "Define and name each session type (Morning, Afternoon, Evening), including start time, end time, and display name. Admins can adjust times here without code changes"),
    ("Payment Terms", "Default invoice due date (e.g. 7 days, 14 days, 30 days). Late payment grace period in days"),
    ("WhatsApp API Configuration", "API key and sender number for Twilio or 360dialog integration. Test mode toggle"),
    ("Email Configuration", "SMTP server, port, username, password, sender email address, sender display name, and branded email template settings"),
    ("Fiscal Year Settings", "Start month of the financial year (e.g. January or July) — used for annual reporting periods"),
    ("Working Hours", "Standard working day hours (used for payroll overtime calculations). Default: 8 hours/day"),
    ("System Access", "Links to role permissions management — defines which roles can access which modules"),
]

for title, desc in settings_areas:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "5.2 User Management", level=2)
add_bullet(doc, "Create, edit, and deactivate staff accounts")
add_bullet(doc, "Assign a role to each user (Super Admin, Admin, Operations, Finance)")
add_bullet(doc, "User profile fields: full name, job title, email address, phone number, role, account status (Active / Inactive), date created")
add_bullet(doc, "Password reset: can be performed by Super Admin or Admin at any time")
add_bullet(doc, "Activity log per user: every action taken by that user is recorded in the audit trail and can be filtered per account")
add_bullet(doc, "Deactivated accounts cannot log in; their records and history are preserved for audit purposes")

add_heading(doc, "5.3 Roles & Permissions", level=2)
add_bullet(doc, "Role-based access control (RBAC) enforced at the API level — middleware rejects unauthorised requests regardless of the UI state")
add_bullet(doc, "Super Admin can adjust which specific actions are permitted per role, within the defined role boundaries")
add_bullet(doc, "Permissions changes are logged in the audit trail with before/after values")
add_bullet(doc, "Any attempt to access a forbidden resource returns a 403 Forbidden response with an appropriate error message displayed in the UI")

add_heading(doc, "5.4 Authentication & Security", level=2)
add_bullet(doc, "Login method: email address + password for all staff roles")
add_bullet(doc, "Two-Factor Authentication (2FA): mandatory for all staff logins. A one-time code is sent via WhatsApp or SMS on every login attempt — the session is not established until the code is verified")
add_bullet(doc, "Session timeout: admin sessions expire after a configurable period of inactivity (default: 30 minutes)")
add_bullet(doc, "Tenant portal: email + password login only; no 2FA required for tenants")
add_bullet(doc, "Password policy: minimum 8 characters, must include a number and a symbol")
add_bullet(doc, "Full Audit Trail: every create, update, delete, approval, and rejection action across the entire platform is logged with: acting user, timestamp, module, record affected, action taken, and before/after field values. The audit trail cannot be edited or deleted by any user including Super Admin.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PHASE 2 — BOOKING SYSTEM
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Phase 2 — Booking System (Pilot)", level=1)

add_heading(doc, "6.1 Product Management", level=2)
add_body(doc, "Before any booking can be created, the building's spaces must be configured as Products in the system. The Product Management module is the single source of truth for all rentable offerings. Staff with Admin or Super Admin access can create, edit, price, and activate or deactivate products.")

add_heading(doc, "Product Types", level=3)
product_types = [
    ("Conference Halls", "Billed per session or custom duration. Can be booked online or by Operations staff."),
    ("Office Spaces", "Monthly lease. Bookings require Admin approval. Tenant onboarding creates the lease record."),
    ("Educational Facility Spaces", "Semester-based lease. Currently the basement is leased to a single university tenant."),
]
for name, desc in product_types:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Per Product Configuration", level=3)
add_bullet(doc, "Name, description, photos (multiple photos uploadable), floor/location")
add_bullet(doc, "Capacity (for halls) or room size/type (for offices)")
add_bullet(doc, "Amenities list: projector, whiteboard, AV system, air conditioning, Wi-Fi, etc. (configurable checkboxes)")
add_bullet(doc, "Base price: per session (halls) or per month (offices)")
add_bullet(doc, "Status: Active (visible and bookable) / Inactive (hidden from public, not bookable)")
add_bullet(doc, "Associated extra services — configured per product type (see below)")
add_bullet(doc, "Dedicated product page on the public website, auto-generated from this data")

add_heading(doc, "Extra Services per Product Type", level=3)
add_bullet(doc, "Conference Halls: Catering packages (Silver / Gold / Platinum), DJ (yes/no + price), Cameraman (yes/no + price), Hall setup / decoration service (optional add-on with configurable price)")
add_bullet(doc, "Office Spaces: Monthly cleaning service (optional add-on), internet add-on, utility billing link")
add_bullet(doc, "Educational Facility: Managed on a custom basis per tenant agreement")

add_heading(doc, "Catering Package Management", level=3)
add_body(doc, "Catering packages are configured in the Central System Settings and are available for selection on any conference hall booking. The three packages are:")

catering_headers = ["Package", "Description", "Typical Inclusions", "Price Model"]
catering_rows = [
    ["Silver",   "Entry-level catering",  "Basic buffet, beverages (water, juice), standard service staff", "Fixed price per package"],
    ["Gold",     "Mid-tier catering",     "Buffet with hot meals, beverages, branded service staff, setup & clearance", "Fixed price per package"],
    ["Platinum", "Premium catering",      "Full gourmet buffet, premium beverages including specialty drinks, dedicated manager, full setup and clearance", "Fixed price per package"],
]
make_table(doc, catering_headers, catering_rows, col_widths=[2.5, 3.5, 7, 3.5])

add_body(doc, "")
add_body(doc, "Each package's included services list and price are fully configurable in System Settings by Super Admin or Admin. DJ and Cameraman are separate add-ons that can be selected independently of the catering package.")

add_heading(doc, "6.2 Conference Hall Booking Flow", level=2)
add_body(doc, "Conference hall bookings follow a structured 4-step approval workflow before a booking is considered confirmed. The Finance Officer's approval is the final gate — no booking is confirmed until Finance approves it.")

add_heading(doc, "4-Step Booking Status Progression", level=3)

status_headers = ["Step", "Status", "Meaning", "Who Acts", "Notification Sent To"]
status_rows = [
    ["1", "Draft",              "Booking submitted online by client, or created by Operations/Admin on behalf of client. Awaiting first review.", "System / Client / Operations / Admin", "Admin (notified of new pending booking)"],
    ["2", "Admin Pending",      "Admin (Building Manager) is reviewing the booking. Must approve or reject.", "Admin", "Admin (action required); Client (acknowledged)"],
    ["3", "Accountant Pending", "Admin has approved the booking. Finance Officer must give final approval after confirming payment is received.", "Finance Officer", "Finance (action required); Client (in progress)"],
    ["4", "Booking Approved",   "Finance Officer has confirmed the booking. Booking is fully confirmed and locked in the calendar.", "Finance Officer", "Client (confirmed + invoice); Admin; Operations"],
]
make_table(doc, status_headers, status_rows, col_widths=[1.2, 3, 4.5, 3.5, 4.3])

add_heading(doc, "Edge Case Statuses", level=3)
edge_statuses = [
    ("Rejected",    "Admin or Finance rejected the booking with a written reason. Client is notified via email and WhatsApp."),
    ("Cancelled",   "An approved booking is cancelled (by client request or building decision). Full refund is processed. Booking is voided in the calendar."),
    ("Rescheduled", "Client requests to change the date/session of an approved booking. The booking reverts to Draft and must go through the full approval chain again."),
    ("Waitlisted",  "The requested slot is already booked. Client is placed on a waiting list and automatically notified if the slot becomes available."),
]
for status, desc in edge_statuses:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{status}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Booking Rules", level=3)
rules = [
    "Bookings can be submitted online by the client (Guest form) or created by Operations/Admin on behalf of a client",
    "Advance booking: up to 3 months in advance from the current date",
    "Priority is given to the client who pays first — Finance will only grant final approval after payment is confirmed",
    "Custom duration: minimum 1 hour, maximum 1 month (in addition to the 3 standard session types)",
    "Recurring bookings: supported (e.g. every Monday morning for 4 weeks); each instance in the series follows the same approval chain",
    "Double-booking prevention: the system blocks a new booking if the slot is already occupied or pending-approved for another client",
    "Cancellation: client notifies building → Admin and Finance are informed → refund is processed → booking is voided",
    "Rescheduling: requires Admin and Finance re-approval; original booking is paused during rescheduling review",
    "Waiting list: system auto-notifies waitlisted clients via WhatsApp and email when a slot opens up",
]
for r in rules:
    add_bullet(doc, r)

add_heading(doc, "Booking Record Fields", level=3)
booking_fields = [
    "Unique booking reference code (e.g. BK-2026-0088)",
    "Client details: full name, company/organisation name, contact phone, email, national ID number",
    "Room selected, session type, date(s), and total duration",
    "Catering package selected: Silver / Gold / Platinum / None",
    "DJ requested: Yes / No (+ price if yes)",
    "Cameraman requested: Yes / No (+ price if yes)",
    "Any other extra services selected",
    "Total price breakdown: base session price + catering + DJ + Cameraman + other extras = grand total",
    "Payment status: Unpaid / Partial / Paid",
    "Status history log: every status change recorded with the acting user, timestamp, and any notes/reason",
    "Communication log: all emails and WhatsApp messages sent to the client linked to this booking",
    "Linked invoice code (auto-generated upon Finance approval)",
    "LPO number (optional — for government, corporate, or NGO clients submitting a Local Purchase Order)",
]
for f in booking_fields:
    add_bullet(doc, f)

add_heading(doc, "Booking Management (Admin Dashboard)", level=3)
add_bullet(doc, "Pending Approvals Queue: separate queues displayed for Admin actions and Finance actions")
add_bullet(doc, "Approve and Reject buttons on each pending booking, with a mandatory notes/reason field for rejections")
add_bullet(doc, "Booking Calendar: month / week / day view across all conference rooms simultaneously, with colour coding per status")
add_bullet(doc, "Operations and Admin can create a booking on behalf of a client from the dashboard (bypasses the online form)")
add_bullet(doc, "Booking Detail Page: shows the full record, status history timeline, communication log, and linked invoice")

add_heading(doc, "6.3 Office & Space Leasing (Long-Term Tenants)", level=2)

add_heading(doc, "Tenant Onboarding", level=3)
add_bullet(doc, "Create a tenant profile: business name, contact person, contact phone, email, national ID, address")
add_bullet(doc, "Document storage: upload lease agreements, KYC documents, and business registration certificates (stored securely in S3)")
add_bullet(doc, "Security Deposit Tracker: record the deposit amount (typically 1 month's rent), date received, and status (Held / Applied to Invoice / Returned)")
add_bullet(doc, "Generate the lease contract from a template (pre-populated with tenant data); support online signing or upload of a scanned signed copy")
add_bullet(doc, "Upon lease creation, generate tenant portal login credentials and email them to the tenant")

add_heading(doc, "Lease Management", level=3)
add_bullet(doc, "Active leases listed with: tenant name, space, floor, lease start date, lease end date, monthly rent, billing cycle, status")
add_bullet(doc, "Automated renewal reminders: email and WhatsApp sent to tenant 10 days before lease expiry; alert sent to Admin")
add_bullet(doc, "Lease renewal workflow: Admin reviews and extends the lease; new contract generated")
add_bullet(doc, "Tenant termination: mark as terminated, generate termination notice, record move-out date, process security deposit return")
add_bullet(doc, "A single tenant can hold multiple office rooms simultaneously (e.g. a company renting 3 rooms on the Ground Floor)")

add_heading(doc, "Office Booking Approval", level=3)
add_body(doc, "Office and educational facility bookings require Admin approval only. The Finance Officer is not part of the approval chain for long-term leases — Finance is involved in the billing and invoicing stage only. Once Admin approves, the lease is created and the tenant is onboarded.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. PHASE 3 — FINANCE & ACCOUNTING
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "7. Phase 3 — Finance & Accounting (Pilot)", level=1)
add_body(doc, "This is the most comprehensive module in the platform. It is designed around professional accounting practices with full double-entry accounting, unique bill and account codes, and complete audit trails for every financial transaction. All financial operations are controlled exclusively by the Finance Officer.")

add_heading(doc, "7.1 Chart of Accounts", level=2)
add_body(doc, "A structured list of all accounts used by Haleelo Tower, each assigned a unique numeric code. The chart of accounts is the backbone of the accounting module — every transaction must be posted to at least two accounts (debit and credit).")

coa_headers = ["Code Range", "Category", "Examples"]
coa_rows = [
    ["1000s", "Assets",      "Cash & bank accounts, Security deposits held, Accounts receivable, Office equipment, Furniture & fixtures"],
    ["2000s", "Liabilities", "Accounts payable (vendor bills), Security deposits owed to tenants, Loans payable"],
    ["3000s", "Revenue",     "Conference hall rental income, Office rental income, Educational facility income, Catering revenue, DJ/Cameraman service revenue, Add-on services revenue"],
    ["4000s", "Expenses",    "Electricity & utilities (water, sewage, garbage), Salaries & wages, Maintenance & repairs, Procurement & supplies, Catering costs, DJ/Cameraman costs, Marketing expenses"],
]
make_table(doc, coa_headers, coa_rows, col_widths=[2.5, 3.5, 10.5])

add_body(doc, "")
add_body(doc, "Finance and Super Admin can create new accounts and assign them to categories at any time. Each account record contains: unique numeric code, account name, category, description, active/inactive status, and current running balance.")

add_heading(doc, "7.2 Account Management (Bank & Mobile Money Accounts)", level=2)
add_body(doc, "The building operates 5 separate accounts that must be tracked independently. Every transaction in the system is tagged to one of these accounts.")

acct_headers = ["Account Name", "Type", "Purpose"]
acct_rows = [
    ["Edahab — Conference Halls",  "Mobile Money", "Receives all conference hall booking payments made via Edahab"],
    ["ZAAD — Conference Halls",    "Mobile Money", "Receives all conference hall booking payments made via ZAAD"],
    ["Edahab — Office Rentals",    "Mobile Money", "Receives all office rent payments made via Edahab"],
    ["ZAAD — Office Rentals",      "Mobile Money", "Receives all office rent payments made via ZAAD"],
    ["Darasalam Bank (Main)",      "Bank Account", "Main building bank account; net profit is transferred here from mobile money accounts after expenses are deducted"],
]
make_table(doc, acct_headers, acct_rows, col_widths=[5, 3, 8.5])

add_body(doc, "")
add_body(doc, "Per Account functionality:")
add_bullet(doc, "Account name, type (bank / mobile money), account number or identifier")
add_bullet(doc, "Current balance (computed in real time from all transactions posted to this account)")
add_bullet(doc, "Transaction history: every payment in and out, with date, reference code, description, and linked invoice or bill code")
add_bullet(doc, "Reconciliation status: Finance can mark a period as reconciled after verifying against the actual bank or mobile money statement")
add_bullet(doc, "Inter-account transfer: record fund movements between any of the 5 accounts (e.g. transferring accumulated balance from ZAAD-Halls to Darasalam Bank), with an auto-generated transfer reference code and a matching journal entry posted automatically")

add_heading(doc, "Account Activity Journal", level=3)
add_body(doc, "A chronological log of every financial event across all accounts, functioning as the General Ledger. Each entry shows: date, description, debit account, credit account, amount, bill/reference code, and posted by. Finance can filter by account, date range, or reference code.")

add_heading(doc, "7.3 Bill Code & Reference System", level=2)
add_body(doc, "Every financial document in the system is assigned a unique, system-generated reference code at the time of creation. These codes appear on all printed and emailed documents and can be searched across the entire platform.")

code_headers = ["Document Type", "Code Format", "Example"]
code_rows = [
    ["Customer Invoice",   "INV-[YEAR]-[SEQ]",  "INV-2026-0042"],
    ["Expense Record",     "EXP-[YEAR]-[SEQ]",  "EXP-2026-0017"],
    ["Vendor Bill",        "VB-[YEAR]-[SEQ]",   "VB-2026-0009"],
    ["Purchase Order",     "PO-[YEAR]-[SEQ]",   "PO-2026-0005"],
    ["Electricity Bill",   "ELEC-[YEAR]-[T#]-[SEQ]", "ELEC-2026-T05-003"],
    ["Payslip",            "PAY-[YEAR]-[EMP#]-[SEQ]", "PAY-2026-E02-012"],
    ["Booking Reference",  "BK-[YEAR]-[SEQ]",   "BK-2026-0088"],
    ["Transfer Reference", "TRF-[YEAR]-[SEQ]",  "TRF-2026-0004"],
]
make_table(doc, code_headers, code_rows, col_widths=[4, 5, 7.5])

add_body(doc, "")
add_body(doc, "Every expense record must include: bill code, expense category (from chart of accounts), payment account (which of the 5 accounts was used), amount, date, description, and an optional receipt or supporting document upload.")

add_heading(doc, "7.4 Customer Invoices & Billing", level=2)

inv_type_headers = ["Invoice Type", "Trigger", "Billing Cycle", "Notes"]
inv_type_rows = [
    ["Office Rent Invoice",           "Auto-generated on billing cycle date per tenant", "Monthly",     "Includes base rent + electricity line item + any add-ons"],
    ["Educational Facility Invoice",  "Manually triggered by Finance",                  "Per semester (3-4 months)", "One invoice per semester per tenant"],
    ["Conference Hall Invoice",       "Auto-generated upon Finance approval of booking", "Per booking", "Includes base session + catering + DJ/Cameraman + extras"],
    ["Electricity Bill",              "After meter reading entry",                       "Monthly",     "Can be standalone or merged as a line item in the rent invoice"],
    ["Manual Invoice",                "Manually created by Finance or Operations",        "Ad hoc",      "For any charge not covered by the above types"],
]
make_table(doc, inv_type_headers, inv_type_rows, col_widths=[4, 4.5, 3.5, 4.5])

add_body(doc, "")
add_body(doc, "Per Invoice fields:")
add_bullet(doc, "Unique bill code (e.g. INV-2026-0042)")
add_bullet(doc, "Client/tenant name and contact details")
add_bullet(doc, "Invoice date and due date")
add_bullet(doc, "Line items: description, quantity, unit price, subtotal")
add_bullet(doc, "LPO number field (for government / corporate / NGO clients)")
add_bullet(doc, "Total amount in USD")
add_bullet(doc, "Payment status: Draft → Sent → Paid / Partial / Overdue")
add_bullet(doc, "Linked payment records (partial and full payment history)")
add_bullet(doc, "Linked account: which of the 5 accounts payment is expected in")
add_bullet(doc, "PDF generation using a branded Haleelo Tower template")
add_bullet(doc, "Email delivery with branded template (via SMTP/SendGrid)")
add_bullet(doc, "WhatsApp delivery option: sends the PDF as an attachment via WhatsApp Business API")

add_heading(doc, "Automated Invoice Behaviour", level=3)
add_bullet(doc, "Office rent invoices are auto-generated on the recurring billing cycle date for each active tenant lease")
add_bullet(doc, "Conference hall invoices are auto-generated at the moment the Finance Officer grants final booking approval")
add_bullet(doc, "Overdue alerts: automated WhatsApp and email reminders are sent when an invoice is 3 days from its due date, on the due date, and at 7 / 14 / 30 days overdue")

add_heading(doc, "7.5 Vendor Bills & Purchase Orders (Procurement)", level=2)
add_body(doc, "The building regularly purchases supplies and services from external vendors. The procurement module tracks orders, bills, and payments in full.")

add_heading(doc, "Vendor Directory", level=3)
add_body(doc, "Finance can create and manage a directory of vendors. Per vendor: business name, contact person, phone, email, category (food supplier, cleaning supplies, equipment, office supplies, event materials, maintenance contractor, etc.), payment terms, and notes.")

add_heading(doc, "Purchase Orders (PO)", level=3)
add_bullet(doc, "Create a PO before ordering: select vendor, add line items (item name, quantity, estimated unit cost), and generate a PO reference code (PO-2026-0005)")
add_bullet(doc, "PO status workflow: Draft → Sent to Vendor → Received → Billed")
add_bullet(doc, "Each PO can be printed or emailed directly to the vendor from the system")

add_heading(doc, "Vendor Bills", level=3)
add_bullet(doc, "Record a vendor bill when an invoice is received from the vendor")
add_bullet(doc, "Assign a bill code: VB-2026-0009")
add_bullet(doc, "Optionally link to an existing PO")
add_bullet(doc, "Line items with descriptions and amounts")
add_bullet(doc, "Expense account code: which account in the chart of accounts this expense is debited to")
add_bullet(doc, "Payment account: which of the 5 building accounts is used to pay this vendor")
add_bullet(doc, "Payment status: Unpaid → Partial → Paid")
add_bullet(doc, "Record payments: date, amount paid, payment method, reference number")

add_heading(doc, "Ad-Hoc Expense Recording", level=3)
add_bullet(doc, "Record non-vendor expenses: utility bills (water, sewage, garbage), maintenance costs, miscellaneous operational costs")
add_bullet(doc, "Required fields: bill code, expense category (chart of accounts), payment account, amount, date, description/notes, optional receipt upload")

add_heading(doc, "7.6 Electricity Bill Management", level=2)
add_body(doc, "Haleelo Tower pays the electricity utility provider and bills each tenant for their individual consumption. The system automates the calculation and billing process based on meter readings.")

add_heading(doc, "Setup (in System Settings)", level=3)
add_bullet(doc, "Electricity rate (USD per kWh): set by Admin or Super Admin in Central Configuration")
add_bullet(doc, "Rate versioning: every rate change is date-stamped. Meter readings always use the rate that was active on the reading date, ensuring historical accuracy")

add_heading(doc, "Meter Reading Entry", level=3)
add_bullet(doc, "Finance or Admin selects a tenant from the list")
add_bullet(doc, "Enters: previous meter reading, current meter reading, reading date")
add_bullet(doc, "System automatically calculates: kWh consumed = current reading − previous reading")
add_bullet(doc, "System automatically calculates: charge = kWh consumed × applicable rate")
add_bullet(doc, "Electricity charge is saved with a unique bill code: ELEC-2026-T05-003")
add_bullet(doc, "The charge can be added to the tenant's next monthly rent invoice as a line item, or issued as a standalone electricity invoice")

elec_headers = ["Field", "Description"]
elec_rows = [
    ["Bill Code",          "Unique electricity bill reference (ELEC-YYYY-T##-###)"],
    ["Tenant",             "Linked tenant record"],
    ["Billing Period",     "Month and year of the reading"],
    ["Previous Reading",   "kWh reading at start of period"],
    ["Current Reading",    "kWh reading at end of period"],
    ["kWh Used",           "Calculated: current − previous"],
    ["Rate Applied",       "USD/kWh rate in effect on the reading date (versioned)"],
    ["Total Charge (USD)", "Calculated: kWh used × rate"],
    ["Invoice Status",     "Draft / Added to Invoice / Invoiced / Paid"],
]
make_table(doc, elec_headers, elec_rows, col_widths=[4, 12.5])

add_heading(doc, "7.7 Payment Recording & Tracking", level=2)
add_bullet(doc, "Record payments against any customer invoice or vendor bill")
add_bullet(doc, "Payment fields: date, amount paid, payment method (Edahab, ZAAD, bank transfer, cheque, cash), reference number (transaction ID from mobile money or bank), which account received or paid")
add_bullet(doc, "Partial payments: fully supported — each payment is recorded individually and the invoice status updates to Partial until fully paid")
add_bullet(doc, "Payment status auto-updates on the invoice: Paid / Partial / Overdue")
add_bullet(doc, "Overpayment handling: excess is recorded as a credit memo against the client's account, to be applied to a future invoice")
add_bullet(doc, "Security Deposit management: record deposit received (amount, date, account received in), track status (Held / Applied to Invoice / Returned), linked to the tenant record")

add_heading(doc, "7.8 Journal Entries & Accounting Ledger", level=2)
add_body(doc, "The accounting module implements full double-entry accounting: every financial event posts both a debit and a credit entry to the chart of accounts, ensuring the books always balance.")

add_bullet(doc, "General Journal: all journal entries in chronological order, auto-generated from invoices, payments, vendor bills, and payroll — plus manual entries where needed")
add_bullet(doc, "Manual Journal Entry: Finance can post manual adjustments, accruals, and corrections. Each manual entry requires a description, debit account, credit account, amount, and supporting note")
add_bullet(doc, "Each journal entry contains: date, description, bill/reference code, debit account + amount, credit account + amount, posted by, timestamp")
add_bullet(doc, "General Ledger view: filter entries by any account code to see a complete history of all transactions affecting that account, with running balance")

add_heading(doc, "7.9 HR & Payroll (Finance-Controlled)", level=2)
add_body(doc, "All payroll operations are managed exclusively by the Finance Officer. There is no employee self-service login. The HR module is a record-keeping and payroll calculation tool used internally by Finance.")

add_heading(doc, "Employee Records", level=3)
add_bullet(doc, "Full profile: full name, role/job title, department (Internal Administration / Maintenance / Cafeteria / Security), start date, contact phone and email, employment type (salaried / daily-rate)")
add_bullet(doc, "Employment contract storage (uploaded PDF, stored in S3)")
add_bullet(doc, "Base salary (for salaried staff) or daily rate (for daily-rate staff) — configurable per employee")
add_bullet(doc, "Employment status: Active / On Leave / Terminated")
add_bullet(doc, "Termination record: last working day, reason, final settlement amount")

add_heading(doc, "Attendance & Leave", level=3)
add_bullet(doc, "Monthly attendance log: Finance records working days, absences, and late arrivals per employee for each calendar month")
add_bullet(doc, "Leave types: Annual Leave, Sick Leave, Unpaid Leave — each tracked with a running balance")
add_bullet(doc, "Leave approval workflow: leave requests (submitted verbally or by paper) are recorded in the system and approved by Admin or Finance")

add_heading(doc, "Overtime & Deductions", level=3)
add_bullet(doc, "Overtime: Finance logs extra hours worked per employee; overtime rate multiplier is configured globally in System Settings (e.g. 1.5× per extra hour) and can be overridden per employee")
add_bullet(doc, "Deductions: unpaid absence days, salary advances repaid, disciplinary deductions — each linked to a unique deduction code and documented with a reason")

add_heading(doc, "Payslip Generation", level=3)
add_bullet(doc, "End-of-month process: Finance runs payroll for the entire workforce")
add_bullet(doc, "System calculates net pay: (base salary or daily rate × days worked) + overtime earnings − deductions")
add_bullet(doc, "PDF payslip generated per employee with all components itemised")
add_bullet(doc, "Payslip delivery: Finance can send via WhatsApp or email directly from the system")
add_bullet(doc, "Payroll expense auto-posted to journal: debit Salary Expense account (chart of accounts), credit the payment account used for salary disbursement")

add_heading(doc, "Payroll Reporting", level=3)
add_bullet(doc, "Monthly payroll summary: total payroll cost broken down by department")
add_bullet(doc, "Individual payroll history: full payslip history per employee")
add_bullet(doc, "Export formats: PDF and Excel")

add_heading(doc, "7.10 Financial Reporting", level=2)
add_body(doc, "All reports are exportable as PDF and Excel. Date range filters apply to all reports.")

add_heading(doc, "Statement Reports", level=3)
stmt_reports = [
    ("Balance Sheet",              "Assets, liabilities, and equity at any selected date"),
    ("Profit & Loss Statement",    "Total revenue vs. total expenses for any period; shows net profit or loss"),
    ("Cash Flow Statement",        "Cash receipts and payments broken down by operating, investing, and financing activities"),
]
for name, desc in stmt_reports:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Audit & Ledger Reports", level=3)
audit_reports = [
    ("Trial Balance",     "All account balances in debit/credit columns at any date — verifies that total debits equal total credits"),
    ("General Ledger",    "All transactions per account, with running balance"),
    ("Journal Audit",     "All journal entries in sequence, with full detail — immutable audit record"),
]
for name, desc in audit_reports:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Partner Reports", level=3)
partner_reports = [
    ("Partner Ledger",      "All transactions for a specific tenant or vendor, in chronological order"),
    ("Aged Receivables",    "Outstanding customer invoices grouped by age: 0-30 days, 31-60 days, 61-90 days, 90+ days"),
    ("Aged Payables",       "Outstanding vendor bills grouped by the same age bands"),
]
for name, desc in partner_reports:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Operational Reports", level=3)
op_reports = [
    ("Revenue Report",           "Total revenue broken down by source: offices, conference halls, educational facility, catering, add-ons"),
    ("Expense Report",           "All expenses by category and account, with supporting reference codes"),
    ("Occupancy Report",         "Occupancy rates per floor, per space type, per time period"),
    ("Bookings Report",          "Detailed list of all bookings with filters: status, room, date range, session type"),
    ("Payment Report",           "All payments — Paid, Pending, and Overdue — with date range filters"),
    ("Electricity Billing Report","Per tenant, per billing period — kWh used, rate applied, total charged, payment status"),
    ("Payroll Report",           "Total payroll cost by department per month; exportable for accounting review"),
]
for name, desc in op_reports:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "Management Dashboard", level=3)
dash_items = [
    "KPI cards: Total Revenue (current month), Active Tenants, Pending Bookings (awaiting action), Overdue Invoices count, Cash & Bank Balances per account",
    "Revenue bar chart: monthly trend over the past 12 months",
    "Revenue breakdown by source (pie or stacked bar)",
    "Outstanding rent and invoice totals",
    "Monthly expenses total and trend",
    "Salary expense totals",
    "Historical comparison: this month vs. the same month in the prior year",
    "Demand analytics: most-booked conference rooms, most popular session types",
    "Quick-action buttons: pending approvals, overdue invoice alerts, upcoming lease renewals",
]
for item in dash_items:
    add_bullet(doc, item)

add_heading(doc, "Automated Reports", level=3)
add_body(doc, "A monthly summary report is automatically generated at the end of each calendar month and emailed to designated management recipients. The report includes a P&L summary, cash position, occupancy snapshot, and outstanding receivables.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. PHASE 4 — MARKETING & PUBLIC WEBSITE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. Phase 4 — Marketing & Public Website (Pilot)", level=1)

add_heading(doc, "8.1 Public Website (halelotower.so)", level=2)
add_body(doc, "The public website is the primary marketing and lead-generation channel for Haleelo Tower. It is fully responsive (mobile and desktop), English only, and presents a professional, high-quality image of the building and its services.")

website_sections = [
    ("Homepage",           "Hero section with building name and high-quality photography, introductory text, and clear call-to-action buttons (Book a Conference Hall, View Office Spaces)"),
    ("Space Showcase",     "Visually rich pages for each category: conference halls, office rooms, educational facility — with professional photos, capacity details, and descriptions"),
    ("How It Works",       "Step-by-step explanation of the booking process so clients understand what to expect"),
    ("Product Pages",      "Dedicated page per conference hall (and per office space type), showing: session times and availability, amenities list, capacity, base pricing, catering package overview (Silver/Gold/Platinum), photo gallery"),
    ("Browse & Filter",    "Conference halls can be filtered by session type, floor, capacity, and amenities"),
    ("Visual Floor Map",   "Interactive floor plan showing each floor's layout; rooms are colour-coded as Available / Booked / Occupied in real time"),
    ("Online Booking Request Form", "Client selects: room, session type, date(s), catering package (Silver/Gold/Platinum/None), DJ (yes/no), Cameraman (yes/no), enters personal/company details, accepts Terms & Conditions, and submits. Form creates a Draft booking in the admin system"),
    ("Booking Acknowledgement", "After form submission: confirmation email sent automatically to the client's email, and a WhatsApp message sent to the client's phone confirming receipt of the booking request and next steps"),
    ("Waiting List",       "If a requested slot is already booked, the client is shown a waiting list option; they enter their contact details and are notified automatically if the slot opens"),
    ("Help & FAQ",         "Answers to common questions about booking, payment, cancellation, and catering"),
    ("Contact Page",       "Building address, phone number, email, and a simple enquiry form"),
]
for name, desc in website_sections:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. PHASE 5 — OTHER MODULES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "9. Phase 5 — Other Modules (Pilot)", level=1)

add_heading(doc, "9.1 Tenant Portal (portal.halelotower.so)", level=2)
add_body(doc, "A secure, branded self-service portal for all building tenants. Access is by email and password (no 2FA required for tenants).")

portal_features = [
    ("In-App Notification Bell", "Tenants receive in-app notifications for: new invoice issued, booking status update, maintenance request update, lease renewal reminder"),
    ("Lease Card",               "Displays space name, floor, lease start date, lease end date, and next payment due date"),
    ("Invoice History",          "Full list of all invoices with status (Paid / Unpaid / Overdue) and a PDF download button per invoice. Electricity bills visible as separate line items"),
    ("Conference Hall Bookings", "Current and past bookings with status; booking details viewable"),
    ("Contract Documents",       "Download the signed lease agreement and other uploaded documents"),
    ("Maintenance Request Form", "Submit a maintenance issue with a text description and optional photo upload"),
    ("Maintenance Tracker",      "View status of all submitted maintenance requests: Open → In Progress → Resolved, with timestamps per status change"),
    ("Account Profile",          "Update contact phone number and email address"),
]
for name, desc in portal_features:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{name}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

add_heading(doc, "9.2 Maintenance Management", level=2)
add_bullet(doc, "Maintenance request inbox in the admin dashboard: displays all tenant-submitted requests with status, priority, and submission date")
add_bullet(doc, "Admin or Operations assigns a staff member or external contractor to each request")
add_bullet(doc, "Status management: Open → In Progress → Resolved (with notes per update)")
add_bullet(doc, "Tenant is automatically notified via WhatsApp and in-portal notification on each status change")
add_bullet(doc, "Work order management: for external repairs, link a vendor/contractor from the vendor directory, record estimated and actual cost, generate a cost record linked to the maintenance request")

add_heading(doc, "9.3 Communications & Announcements", level=2)
add_bullet(doc, "Broadcast Announcements: Admin or Operations sends a message to all active tenants simultaneously via email and WhatsApp")
add_bullet(doc, "Automated WhatsApp reminders: lease renewal reminder (10 days before expiry), rent due reminder (3 days before due date), overdue payment alert (at due date, and at 7/14/30 days overdue)")
add_bullet(doc, "Automated emails: booking confirmation, booking approval/rejection notification, invoice issued, maintenance request status update")
add_bullet(doc, "Branded email templates: all emails use the Haleelo Tower logo, building colours, and a professional footer with contact details")
add_bullet(doc, "All sent messages are logged per tenant in the communication log (visible in the tenant profile and in the linked booking or invoice record)")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 10. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "10. Technology Stack", level=1)

tech_headers = ["Layer", "Technology", "Purpose"]
tech_rows = [
    ["Frontend",          "Next.js (React)",         "Admin dashboard, tenant portal, and public website — all three subdomains"],
    ["Backend / API",     "Laravel (PHP)",            "RESTful API server; business logic, authentication, authorisation, and all data processing"],
    ["Database",          "PostgreSQL",               "Primary relational database for all transactional and accounting data"],
    ["File Storage",      "AWS S3",                   "Document and photo storage: lease agreements, KYC documents, payslips, receipts, product photos"],
    ["Messaging",         "WhatsApp Business API",    "Automated WhatsApp notifications via Twilio or 360dialog integration"],
    ["Email",             "SMTP / SendGrid",          "Transactional email delivery with branded HTML templates"],
    ["Authentication",    "JWT + 2FA",                "Stateless token-based auth for admin; OTP sent via WhatsApp/SMS for 2FA"],
    ["PDF Generation",    "Server-side PDF library",  "Invoice PDFs, payslips, lease contracts, financial reports"],
    ["Hosting",           "VPS",                      "Local-first testing environment; production on VPS with SSL (Let's Encrypt) for all 3 subdomains"],
    ["Currency",          "USD only",                 "All monetary values stored and displayed in US Dollars"],
    ["Language",          "English only",             "All UI, documents, and communications in English"],
]
make_table(doc, tech_headers, tech_rows, col_widths=[3.5, 4, 9])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 11. DEPLOYMENT & INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "11. Deployment & Infrastructure", level=1)

add_heading(doc, "11.1 Subdomain Structure", level=2)
subdomain_headers = ["Subdomain", "Application", "Deployment"]
subdomain_rows = [
    ["halelotower.so",        "Next.js public website",                "VPS — Nginx reverse proxy, SSL via Let's Encrypt"],
    ["portal.halelotower.so", "Next.js tenant portal",                 "VPS — Nginx reverse proxy, SSL via Let's Encrypt"],
    ["admin.halelotower.so",  "Next.js admin dashboard + Laravel API", "VPS — Nginx reverse proxy, SSL via Let's Encrypt"],
]
make_table(doc, subdomain_headers, subdomain_rows, col_widths=[5, 5, 6.5])

add_heading(doc, "11.2 Development & Testing Approach", level=2)
add_bullet(doc, "Local-first development: all development and internal testing conducted in a local environment before deployment to the VPS")
add_bullet(doc, "Staging environment on VPS: a staging subdomain for client review and UAT before each phase goes live")
add_bullet(doc, "Production deployment: only after client sign-off on each phase")
add_bullet(doc, "Database backups: automated daily database backups with 30-day retention")
add_bullet(doc, "Zero-downtime deployments using rolling restart strategy on VPS")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 12. PHASED DELIVERY PLAN
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "12. Phased Delivery Plan", level=1)
add_body(doc, "The implementation is structured into 5 development phases followed by a testing and go-live phase. Estimated durations reflect the increased scope confirmed in the June 2026 discovery sessions, particularly the full accounting module in Phase 3.")

delivery_headers = ["Phase", "Focus Area", "Key Deliverables", "Estimated Duration"]
delivery_rows = [
    ["Phase 1", "User Management, Roles, Settings, Authentication",
     "Staff accounts, RBAC enforcement, 2FA, audit trail, central system settings panel, electricity rate config, catering package config",
     "1 week"],
    ["Phase 2", "Product Management & Booking System",
     "Product pages, catering packages, 4-step conference hall booking flow, office leasing workflow, booking calendar, waiting list, online booking form integration",
     "1.5 weeks"],
    ["Phase 3", "Finance & Accounting",
     "Chart of accounts, 5-account management, bill code system, customer invoicing, vendor bills, procurement, electricity billing, payment recording, double-entry journal, payroll, all financial reports, management dashboard",
     "3 weeks"],
    ["Phase 4", "Public Website",
     "Marketing pages, product pages, interactive floor map, online booking form, FAQ, booking acknowledgements",
     "1 week"],
    ["Phase 5", "Tenant Portal, Maintenance, Communications",
     "Tenant portal with invoice history, maintenance tracker, contract downloads; maintenance management module; automated WhatsApp and email notifications",
     "1 week"],
    ["Final", "End-to-End Testing, Training & Go-Live",
     "Full regression testing across all modules, staff training sessions per role, go-live support and monitoring",
     "0.5 weeks"],
    ["TOTAL", "", "", "~8 weeks"],
]
make_table(doc, delivery_headers, delivery_rows, col_widths=[2, 4, 9, 2.5])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 13. TRAINING & HANDOVER
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "13. Training & Handover", level=1)
add_body(doc, "Structured training sessions will be delivered per role group before the go-live date. All training will be conducted in English, on-site in the Haleelo Tower building.")

training_headers = ["Role", "Training Focus", "Format"]
training_rows = [
    ["Super Admin",              "Full system walkthrough: settings configuration, user management, audit trail review, system maintenance, backup verification",
     "1-on-1 session, 3 hours"],
    ["Admin / Building Manager", "Booking approvals and calendar, tenant onboarding, lease management, maintenance assignments, broadcast communications, electricity meter readings",
     "Group session, 2 hours"],
    ["Operations / Receptionist","Creating bookings on behalf of clients, tenant onboarding steps, invoice generation, using the booking calendar, maintenance request management",
     "Group session, 2 hours"],
    ["Finance / Accountant",     "Full accounting module: chart of accounts setup, invoicing, payment recording, vendor bills, electricity billing, journal entries, payroll processing, account code system, bill code system, all financial reports, trial balance, management dashboard, final booking approval workflow",
     "Dedicated session, 4 hours"],
]
make_table(doc, training_headers, training_rows, col_widths=[4, 8.5, 4])

add_body(doc, "")
add_body(doc, "A written user guide will be provided for each role, tailored to their specific modules and workflows. Video recordings of the training sessions will be shared for future reference and onboarding of new staff.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 14. POST-LAUNCH SUPPORT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "14. Post-Launch Support", level=1)
add_body(doc, "Following the go-live date, the development team will provide a post-launch support period to address any issues, bugs, or usability concerns that arise in real-world use.")

support_items = [
    ("Hypercare Period (Weeks 1-2 post-launch)", "Daily availability for bug fixes and urgent issues; priority response within 4 hours"),
    ("Standard Support (Month 1-3 post-launch)", "Response to bug reports within 1 business day; minor feature adjustments as needed"),
    ("Change Requests",                           "Any new features or significant changes to existing features discovered post-launch are scoped as separate change requests and quoted accordingly"),
    ("Data Migration",                            "If existing records (tenants, lease data, transaction history) need to be imported, this will be scoped separately before Phase 1 development begins"),
    ("Ongoing Hosting Support",                   "VPS configuration, SSL renewal, and database backup monitoring are included in the support scope"),
]
for title, desc in support_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 15. RECOMMENDATIONS FOR REVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "15. Recommendations for Review", level=1)
add_body(doc, "The following features are recommended by the development team based on professional platform standards and common requirements for buildings of this type. These were not explicitly specified in the discovery sessions and are not included in the current development scope. The client should review each item and confirm which to include before development begins. Including additional items will affect the project timeline and cost.")

recommendations = [
    ("1. Petty Cash Management",
     "Track small day-to-day cash purchases separately from the main 5 accounts. Maintains a running petty cash balance, records each spend with a code and description, and includes a replenishment workflow when the float runs low."),
    ("2. Bank Reconciliation Tool",
     "A dedicated tool that lets Finance upload or manually enter the bank/mobile money statement and match it line-by-line against the system's recorded transactions. Highlights discrepancies — missing entries, duplicates, or mismatched amounts — for investigation."),
    ("3. Late Payment Penalties",
     "The system automatically calculates and applies a configurable late payment fee (fixed amount or percentage) to overdue invoices after the grace period expires. Rate is configured in System Settings."),
    ("4. Receipt / Payment Vouchers",
     "Generate a formal payment receipt document (printable PDF) for every payment received from a tenant or client. The receipt prominently displays the bill code, amount, payment method, and date. Useful for clients who require physical receipts."),
    ("5. Inter-Account Transfer Records",
     "A dedicated workflow to formally record fund transfers between the 5 accounts (e.g. moving weekly revenue from ZAAD-Halls to Darasalam Bank). Generates a transfer reference code (TRF-2026-0004) and automatically posts the correct journal entries (debit receiving account, credit sending account)."),
    ("6. Budget Planning vs. Actuals",
     "Finance sets a monthly or annual budget for each expense category. The system tracks actual spending against the budget in real time and sends an alert when a category approaches 80% or exceeds 100% of its budget."),
    ("7. Accounts Receivable Aging Alerts",
     "Automated escalating alerts as customer invoices age: WhatsApp reminder at 30 days overdue, escalation to Admin at 60 days, management dashboard alert at 90 days overdue. Reduces manual chasing effort."),
    ("8. Fiscal Year Lock",
     "Finance or Super Admin can lock a completed financial period (e.g. end of fiscal year) so no edits, deletions, or backdated entries can be made to historical transactions. Ensures audit integrity."),
    ("9. Catering Expense Tracking Per Event",
     "When Finance pays vendor bills for catering materials for a specific confirmed conference hall booking, those expenses are linked directly to the booking record. This enables per-event profitability calculation: (booking revenue) − (direct catering costs) = event margin."),
    ("10. Asset Register",
     "Track the building's major assets: furniture, AV equipment, generator, air conditioning units, etc. Per asset: purchase date, purchase value, current estimated value, and depreciation schedule. Useful for insurance, audits, and replacement planning."),
    ("11. Credit Notes",
     "Issue a formal credit note against any customer invoice for partial or full refunds. The credit note automatically reduces the tenant's outstanding balance and posts the correct reversing journal entry. Linked to the original invoice by bill code."),
    ("12. Multi-Currency Display (Future-Proofing)",
     "Even though the building operates in USD only, design the database schema to support a second currency field from the outset. This means that if Somali Shilling billing or reporting is needed in future, the schema change is minimal and no data migration is required."),
    ("13. Recurring Expense Templates",
     "Set up recurring expense entries that are auto-generated each month (e.g. electricity utility bill, internet provider bill, garbage collection fee). Finance reviews and confirms or adjusts the auto-generated record before posting. Saves significant monthly data entry time."),
    ("14. Tax / VAT Configuration",
     "Add a configurable tax rate field in System Settings and on invoices — even if the current rate is 0%. If taxation requirements change in Somalia in the future, the system is already prepared and Finance can simply update the rate without any code changes."),
]

for title, desc in recommendations:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 16. COMPLETE FEATURE REGISTER
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "16. Complete Feature Register", level=1)
add_body(doc, "The table below is the full list of confirmed platform features as of v2.0. All features listed are included in the development scope unless marked as Future. Phase number indicates the delivery phase in which the feature is built.")

feat_headers = ["#", "Feature", "Description", "Phase", "Status"]
features = [
    # Phase 1
    (1,  "Central System Settings",           "Unified configuration panel in admin dashboard for all platform settings", "1", "Pilot"),
    (2,  "General Settings",                  "Building name, logo, contact details, address", "1", "Pilot"),
    (3,  "Electricity Rate Configuration",    "Set and version USD/kWh electricity rate; historical rates preserved", "1", "Pilot"),
    (4,  "Catering Package Configuration",    "Create/edit Silver, Gold, Platinum packages with services and prices", "1", "Pilot"),
    (5,  "Session Time Configuration",        "Configure morning, afternoon, evening session names and times", "1", "Pilot"),
    (6,  "Payment Terms Configuration",       "Default invoice due dates, late payment grace period", "1", "Pilot"),
    (7,  "WhatsApp API Configuration",        "API key and sender number (Twilio / 360dialog)", "1", "Pilot"),
    (8,  "Email / SMTP Configuration",        "SMTP settings, sender name, branded templates", "1", "Pilot"),
    (9,  "Fiscal Year Settings",              "Configure start month of financial year", "1", "Pilot"),
    (10, "Working Hours Configuration",       "Standard working day hours for payroll overtime", "1", "Pilot"),
    (11, "User Account Management",           "Create, edit, deactivate staff accounts; user profiles", "1", "Pilot"),
    (12, "Role-Based Access Control (RBAC)",  "4 roles enforced at API level: Super Admin, Admin, Operations, Finance", "1", "Pilot"),
    (13, "Email + Password Authentication",   "Secure login for all staff roles", "1", "Pilot"),
    (14, "Two-Factor Authentication (2FA)",   "Mandatory OTP via WhatsApp/SMS for all staff logins", "1", "Pilot"),
    (15, "Session Timeout",                   "Automatic logout after configurable inactivity period", "1", "Pilot"),
    (16, "Audit Trail",                       "Immutable log of all create, update, delete, approval actions with before/after values", "1", "Pilot"),
    (17, "Password Reset",                    "Admin/Super Admin can reset any staff account password", "1", "Pilot"),
    # Phase 2
    (18, "Product Management Module",         "Create, edit, price, activate/deactivate building products", "2", "Pilot"),
    (19, "Conference Hall Products",          "Configure halls with capacity, amenities, photos, pricing, sessions", "2", "Pilot"),
    (20, "Office Space Products",             "Configure office rooms with floor, size, photos, monthly pricing", "2", "Pilot"),
    (21, "Educational Facility Product",      "Configure basement educational facility product and pricing", "2", "Pilot"),
    (22, "Extra Services per Product",        "Attach catering, cleaning, DJ, Cameraman, decoration add-ons", "2", "Pilot"),
    (23, "Catering Package Selection",        "Silver/Gold/Platinum selectable at conference hall booking time", "2", "Pilot"),
    (24, "DJ Add-On",                         "Yes/No + price, selectable at booking time for conference halls", "2", "Pilot"),
    (25, "Cameraman Add-On",                  "Yes/No + price, selectable at booking time for conference halls", "2", "Pilot"),
    (26, "4-Step Booking Status Flow",        "Draft → Admin Pending → Accountant Pending → Booking Approved", "2", "Pilot"),
    (27, "Admin Approval Step",               "Admin reviews and approves/rejects conference hall bookings (first gate)", "2", "Pilot"),
    (28, "Finance Approval Step",             "Finance Officer final approval of conference hall bookings (second gate)", "2", "Pilot"),
    (29, "Rejection with Reason",             "Admin or Finance can reject with a written reason; client notified", "2", "Pilot"),
    (30, "Booking Reference Code",            "Unique code per booking (BK-2026-XXXX)", "2", "Pilot"),
    (31, "Booking Calendar",                  "Month/week/day view across all rooms with colour-coded status", "2", "Pilot"),
    (32, "Create Booking on Behalf",          "Operations or Admin creates booking for client from dashboard", "2", "Pilot"),
    (33, "Recurring Bookings",                "Support for weekly/regular recurring conference hall bookings", "2", "Pilot"),
    (34, "Custom Duration Bookings",          "Min 1 hour to max 1 month, in addition to standard sessions", "2", "Pilot"),
    (35, "Waiting List",                      "Auto-notify waitlisted clients when a requested slot opens", "2", "Pilot"),
    (36, "Booking Cancellation Workflow",     "Client notifies building; Admin/Finance informed; refund processed; slot freed", "2", "Pilot"),
    (37, "Rescheduling Workflow",             "Approved booking can be rescheduled; triggers full re-approval chain", "2", "Pilot"),
    (38, "Booking Status History Log",        "Per-booking log of every status change with user, timestamp, and notes", "2", "Pilot"),
    (39, "Booking Communication Log",         "All emails and WhatsApp messages linked to the booking record", "2", "Pilot"),
    (40, "Pending Approvals Queue",           "Separate queues for Admin actions and Finance actions on the dashboard", "2", "Pilot"),
    (41, "Tenant Profile",                    "Business name, contacts, national ID, documents, security deposit", "2", "Pilot"),
    (42, "Tenant Document Storage",           "Upload and store KYC documents, lease agreements, business registration", "2", "Pilot"),
    (43, "Security Deposit Tracker",          "Record deposit, status (Held/Applied/Returned), linked to tenant", "2", "Pilot"),
    (44, "Lease Contract Generation",         "Auto-populate lease template with tenant data; upload or sign online", "2", "Pilot"),
    (45, "Lease Renewal Workflow",            "Automated reminders 10 days before expiry; Admin extends lease", "2", "Pilot"),
    (46, "Tenant Termination",                "Mark terminated, generate notice, record move-out, process deposit return", "2", "Pilot"),
    (47, "Tenant Portal Access Provisioning", "Auto-generate and email portal login credentials when lease created", "2", "Pilot"),
    (48, "Multi-Room Tenancy",                "Single tenant can hold multiple office rooms simultaneously", "2", "Pilot"),
    # Phase 3
    (49, "Chart of Accounts",                 "Structured account list with unique numeric codes (1000s–4000s categories)", "3", "Pilot"),
    (50, "Account Management (5 Accounts)",   "Track Edahab×2, ZAAD×2, Darasalam Bank separately with balances", "3", "Pilot"),
    (51, "Account Transaction History",       "Every payment in/out per account with date, reference, description", "3", "Pilot"),
    (52, "Account Reconciliation",            "Finance marks periods as reconciled per account", "3", "Pilot"),
    (53, "Inter-Account Transfer Recording",  "Record fund movements between the 5 accounts with auto journal entry", "3", "Pilot"),
    (54, "Bill Code & Reference System",      "Unique system-generated codes for every financial document", "3", "Pilot"),
    (55, "Office Rent Invoicing",             "Auto-generated monthly invoices per tenant with line items", "3", "Pilot"),
    (56, "Conference Hall Invoicing",         "Auto-generated upon Finance booking approval with full price breakdown", "3", "Pilot"),
    (57, "Educational Facility Invoicing",    "Semester invoices generated by Finance", "3", "Pilot"),
    (58, "Electricity Bill Invoicing",        "Standalone or merged line item in rent invoice", "3", "Pilot"),
    (59, "Manual Invoice Creation",           "One-off invoices for ad-hoc charges", "3", "Pilot"),
    (60, "LPO Number Field",                  "Support for Local Purchase Order numbers on invoices", "3", "Pilot"),
    (61, "Invoice PDF Generation",            "Branded PDF with all line items, due date, and bill code", "3", "Pilot"),
    (62, "Invoice Email Delivery",            "Send invoice PDF via SMTP/SendGrid with branded template", "3", "Pilot"),
    (63, "Invoice WhatsApp Delivery",         "Send invoice PDF via WhatsApp Business API", "3", "Pilot"),
    (64, "Overdue Invoice Alerts",            "Automated WhatsApp + email reminders at due date and overdue intervals", "3", "Pilot"),
    (65, "Vendor Directory",                  "Create and manage vendors with categories and payment terms", "3", "Pilot"),
    (66, "Purchase Orders (PO)",              "Create POs with items and estimated costs; status workflow", "3", "Pilot"),
    (67, "Vendor Bills",                      "Record vendor invoices with expense and payment account codes", "3", "Pilot"),
    (68, "Ad-Hoc Expense Recording",          "Record utility bills, maintenance costs with receipt upload", "3", "Pilot"),
    (69, "Electricity Rate Versioning",        "Date-stamped rate history; old readings use old rates automatically", "3", "Pilot"),
    (70, "Meter Reading Entry & Calculation", "Admin enters readings; system calculates kWh used and USD charge", "3", "Pilot"),
    (71, "Electricity Bill Record",           "Full record per reading with unique bill code", "3", "Pilot"),
    (72, "Payment Recording",                 "Record payments against invoices/bills with full payment details", "3", "Pilot"),
    (73, "Partial Payment Support",           "Multiple partial payments tracked per invoice/bill", "3", "Pilot"),
    (74, "Overpayment / Credit Handling",     "Excess payments recorded as credits against future invoices", "3", "Pilot"),
    (75, "Double-Entry Accounting",           "Every transaction posts a debit and credit to chart of accounts", "3", "Pilot"),
    (76, "General Journal",                   "Chronological journal of all financial events across all accounts", "3", "Pilot"),
    (77, "Manual Journal Entry",              "Finance can post adjustments, accruals, and corrections", "3", "Pilot"),
    (78, "General Ledger View",               "Filter journal by any account; running balance shown", "3", "Pilot"),
    (79, "Employee Records",                  "Full staff profiles: name, role, department, contract, salary, status", "3", "Pilot"),
    (80, "Attendance & Leave Tracking",       "Monthly attendance log, leave types and balances per employee", "3", "Pilot"),
    (81, "Overtime & Deduction Management",   "Log overtime hours and salary deductions per employee per month", "3", "Pilot"),
    (82, "Payslip Generation (PDF)",          "Auto-calculated net pay; branded PDF payslip per employee", "3", "Pilot"),
    (83, "Payslip Delivery",                  "Finance sends payslip via WhatsApp or email from dashboard", "3", "Pilot"),
    (84, "Payroll Journal Posting",           "Payroll expense auto-posted to journal: debit Salary, credit Payment Acct", "3", "Pilot"),
    (85, "Balance Sheet Report",              "Assets, liabilities, equity at any date — PDF and Excel export", "3", "Pilot"),
    (86, "Profit & Loss Report",              "Revenue vs. expenses for any period — PDF and Excel export", "3", "Pilot"),
    (87, "Cash Flow Statement",               "Cash movements in/out by activity type — PDF and Excel export", "3", "Pilot"),
    (88, "Trial Balance Report",              "All account balances debit/credit at any date; verifies books balance", "3", "Pilot"),
    (89, "General Ledger Report",             "All transactions per account with running balance — exportable", "3", "Pilot"),
    (90, "Journal Audit Report",              "All journal entries in sequence — immutable audit report", "3", "Pilot"),
    (91, "Partner Ledger",                    "All transactions per tenant or vendor — exportable", "3", "Pilot"),
    (92, "Aged Receivables Report",           "Outstanding invoices grouped by age bands (0-30, 31-60, 61-90, 90+)", "3", "Pilot"),
    (93, "Aged Payables Report",              "Outstanding vendor bills grouped by age bands", "3", "Pilot"),
    (94, "Revenue Report",                    "Revenue by source: halls, offices, educational, add-ons", "3", "Pilot"),
    (95, "Expense Report",                    "All expenses by category and account code", "3", "Pilot"),
    (96, "Occupancy Report",                  "Occupancy rates per floor, space type, period", "3", "Pilot"),
    (97, "Bookings Report",                   "Detailed filterable list of all bookings", "3", "Pilot"),
    (98, "Payment Report",                    "Paid / Pending / Overdue payments with date filters", "3", "Pilot"),
    (99, "Electricity Billing Report",        "Per tenant, per period — kWh, rate, charge, status", "3", "Pilot"),
    (100, "Payroll Report",                   "Payroll cost by department per month — PDF and Excel", "3", "Pilot"),
    (101, "Management Dashboard",             "KPI cards, revenue charts, occupancy snapshot, alerts", "3", "Pilot"),
    (102, "Automated Monthly Report Email",   "Monthly summary auto-emailed to management at month end", "3", "Pilot"),
    # Phase 4
    (103, "Public Website Homepage",          "Hero section, building intro, CTA buttons", "4", "Pilot"),
    (104, "Space Showcase Pages",             "Category pages for halls, offices, educational facility with photos", "4", "Pilot"),
    (105, "How It Works Section",             "Step-by-step booking process explanation", "4", "Pilot"),
    (106, "Product Pages (Per Hall/Office)",  "Individual page per product with sessions, amenities, pricing, gallery", "4", "Pilot"),
    (107, "Browse & Filter",                  "Filter conference halls by session, floor, capacity, amenities", "4", "Pilot"),
    (108, "Interactive Floor Map",            "Visual floor plan with real-time availability colour coding", "4", "Pilot"),
    (109, "Online Booking Request Form",      "Client selects room, session, catering, DJ, Cameraman; submits; creates Draft booking", "4", "Pilot"),
    (110, "Booking Acknowledgement",          "Auto confirmation email + WhatsApp to client on form submission", "4", "Pilot"),
    (111, "Waiting List (Public)",            "Client joins waiting list if slot is booked; auto-notified on availability", "4", "Pilot"),
    (112, "Help & FAQ Page",                  "Common questions about booking, payment, cancellation, catering", "4", "Pilot"),
    (113, "Responsive Design",                "Fully mobile and desktop responsive across all public pages", "4", "Pilot"),
    # Phase 5
    (114, "Tenant Portal Login",              "Secure email + password login for tenants (no 2FA)", "5", "Pilot"),
    (115, "Tenant In-App Notifications",      "Bell icon with notifications for invoices, bookings, maintenance, renewals", "5", "Pilot"),
    (116, "Tenant Lease Card",                "Space name, floor, lease dates, next payment due", "5", "Pilot"),
    (117, "Tenant Invoice History",           "All invoices with status and PDF download; electricity visible separately", "5", "Pilot"),
    (118, "Tenant Booking History",           "Current and past conference room bookings with status detail", "5", "Pilot"),
    (119, "Tenant Document Download",         "Download signed lease agreement and uploaded documents", "5", "Pilot"),
    (120, "Maintenance Request Submission",   "Tenant submits request with description and optional photo", "5", "Pilot"),
    (121, "Maintenance Request Tracker",      "Tenant views status: Open → In Progress → Resolved with timestamps", "5", "Pilot"),
    (122, "Tenant Profile Update",            "Tenant can update contact phone and email", "5", "Pilot"),
    (123, "Maintenance Request Inbox",        "Admin dashboard inbox with all requests, status, and assignment", "5", "Pilot"),
    (124, "Maintenance Staff Assignment",     "Admin assigns a staff member or contractor per request", "5", "Pilot"),
    (125, "Maintenance Status Management",    "Admin updates status; tenant auto-notified on each change", "5", "Pilot"),
    (126, "Maintenance Work Orders",          "Link external vendor for repairs; record cost against request", "5", "Pilot"),
    (127, "Broadcast Announcements",          "Admin sends message to all tenants via email + WhatsApp simultaneously", "5", "Pilot"),
    (128, "Automated Lease Renewal Reminder", "WhatsApp + email to tenant 10 days before expiry; Admin alert", "5", "Pilot"),
    (129, "Automated Rent Due Reminder",      "WhatsApp + email 3 days before invoice due date", "5", "Pilot"),
    (130, "Automated Overdue Payment Alert",  "Escalating WhatsApp + email at due date, 7, 14, and 30 days overdue", "5", "Pilot"),
    (131, "Branded Email Templates",          "Haleelo Tower logo, colours, and professional footer on all outgoing emails", "5", "Pilot"),
    (132, "Communication Log",                "All messages sent linked to tenant, booking, or invoice record", "5", "Pilot"),
]

make_table(doc, feat_headers, features, col_widths=[0.8, 4, 7.8, 1.5, 1.9])

# ── Save ──────────────────────────────────────────────────────────────────────
save_path = "/home/user/haleelo-tower/Haleelo Tower — Platform Implementation Plan v2.0.docx"
doc.save(save_path)
print(f"Saved: {save_path}")
